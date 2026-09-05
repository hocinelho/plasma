"""Tests for meeting recording, summarizing and Word export.

Nothing here touches the microphone or an LLM — the recorder thread is driven
through its transcription step directly.
"""
from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from backend.modules.meeting import export_docx, recorder, summarizer  # noqa: E402
from backend.skills import meeting_notes  # noqa: E402


def _state(segments=None, title="Test meeting"):
    return recorder.MeetingState(
        meeting_id="20260811-140000",
        title=title,
        started_at=datetime(2026, 8, 11, 14, 0, 0),
        ended_at=datetime(2026, 8, 11, 14, 30, 0),
        segments=segments if segments is not None else [
            {"ts": "2026-08-11T14:00:05", "text": "We need to decide the rollout."},
            {"ts": "2026-08-11T14:02:00", "text": "Hocine prepares the measurements."},
        ],
    )


# ── skill ────────────────────────────────────────────────────────────────
def test_self_test_passes():
    assert meeting_notes.self_test() is True


def test_status_when_idle():
    reply = meeting_notes.run({"utterance": "meeting status"})
    assert "not recording" in reply.lower()


def test_stopping_without_a_meeting_is_not_an_error():
    reply = meeting_notes.run({"utterance": "stop the meeting"})
    assert "no meeting" in reply.lower()


@pytest.mark.parametrize("utterance,title", [
    ("start meeting notes about the fibre rollout", "the fibre rollout"),
    ("start meeting notes on budget planning", "budget planning"),
    ("protokoll starten für die Netzplanung", "die Netzplanung"),
    ("start meeting notes", ""),
])
def test_title_extraction(utterance, title):
    assert meeting_notes._title_from(utterance) == title


# ── transcript state ─────────────────────────────────────────────────────
@pytest.mark.parametrize("utterance", [
    "start meeting notes",
    "Start the meeting notes.",
    "record this meeting",
    "Hey Plasma, record the meeting",
    "take meeting notes",
    "start recording the meeting",   # open_app owns the generic "start "
    "start recording",
    "stop the meeting",
    "stop recording",
    "end the meeting",
    "meeting status",
    "Protokoll starten",
    "Meeting aufnehmen",
    "meeting beenden",
])
def test_spoken_commands_route_to_this_skill(utterance):
    """Trigger matching is longest-wins, so generic triggers in other skills
    can steal a phrase — 'start recording the meeting' went to open_app."""
    from backend.modules.skills.registry import SkillRegistry

    registry = SkillRegistry()
    registry.load_all()
    skill = registry.find_by_trigger(utterance)
    assert skill is not None and skill.name == "meeting_notes", utterance


@pytest.mark.parametrize("utterance,expected", [
    ("open spotify", "open_app"),
    ("launch chrome", "open_app"),
])
def test_other_skills_are_not_stolen_back(utterance, expected):
    from backend.modules.skills.registry import SkillRegistry

    registry = SkillRegistry()
    registry.load_all()
    skill = registry.find_by_trigger(utterance)
    assert skill is not None and skill.name == expected


class _AliveThread:
    """Stand-in for the recorder thread, so no microphone is needed."""

    def is_alive(self):
        return True


@pytest.fixture
def _pretend_recording():
    original = recorder.recorder._thread
    recorder.recorder._thread = _AliveThread()
    yield
    recorder.recorder._thread = original


def test_hands_free_triggers_are_muted_during_a_meeting(_pretend_recording):
    """Wake word and clap fire on whatever the room says.

    During a meeting that made Plasma talk over the meeting and send the
    room's words to the LLM as commands.
    """
    from backend.modules.voice.wake_monitor import _hands_free_suppressed

    assert _hands_free_suppressed() is True


def test_hands_free_triggers_work_again_once_the_meeting_ends():
    from backend.modules.voice.wake_monitor import _hands_free_suppressed

    assert recorder.recorder.is_recording is False
    assert _hands_free_suppressed() is False


def test_suppression_never_raises_if_the_meeting_module_is_unavailable(monkeypatch):
    """The wake loop must not die because of a meeting import problem."""
    import builtins

    from backend.modules.voice import wake_monitor

    real_import = builtins.__import__

    def boom(name, *a, **kw):
        if "meeting" in name:
            raise ImportError("simulated")
        return real_import(name, *a, **kw)

    monkeypatch.setattr(builtins, "__import__", boom)
    assert wake_monitor._hands_free_suppressed() is False


def test_transcript_joins_segments():
    assert _state().transcript.splitlines() == [
        "We need to decide the rollout.",
        "Hocine prepares the measurements.",
    ]


def test_duration_is_computed():
    assert _state().duration_min == pytest.approx(30.0)


def test_persisted_meeting_can_be_reloaded(tmp_path, monkeypatch):
    """An hour of meeting must survive a crash — segments are written as they go."""
    monkeypatch.setattr(recorder, "MEETINGS_DIR", tmp_path)
    segments = [
        {"ts": "2026-08-11T14:00:05", "text": "first"},
        {"ts": "2026-08-11T14:00:35", "text": "second"},
    ]
    path = tmp_path / "20260811-140000.jsonl"
    path.write_text("".join(json.dumps(s) + "\n" for s in segments), encoding="utf-8")

    loaded = recorder.load("20260811-140000")
    assert loaded is not None
    assert [s["text"] for s in loaded.segments] == ["first", "second"]
    assert recorder.load_latest().meeting_id == "20260811-140000"


def test_corrupt_transcript_lines_are_skipped(tmp_path, monkeypatch):
    monkeypatch.setattr(recorder, "MEETINGS_DIR", tmp_path)
    (tmp_path / "20260811-140000.jsonl").write_text(
        '{"ts":"2026-08-11T14:00:05","text":"good"}\nnot json at all\n\n',
        encoding="utf-8",
    )
    loaded = recorder.load("20260811-140000")
    assert [s["text"] for s in loaded.segments] == ["good"]


# ── summarizer ───────────────────────────────────────────────────────────
def test_json_is_extracted_from_a_fenced_reply():
    raw = ('Sure!\n```json\n{"summary":"S","key_points":["a"],"decisions":[],'
           '"actions":[],"open_questions":[]}\n```')
    assert summarizer._extract_json(raw)["summary"] == "S"


def test_json_is_extracted_when_wrapped_in_prose():
    raw = 'Here you go: {"summary":"S","key_points":[]} — hope that helps!'
    assert summarizer._extract_json(raw)["summary"] == "S"


def test_unparseable_reply_returns_none():
    assert summarizer._extract_json("I could not do that.") is None
    assert summarizer._extract_json("") is None


def test_empty_transcript_is_reported_not_invented():
    result = summarizer.summarize("")
    assert result["degraded"] is True
    assert result["decisions"] == []


def test_fallback_keeps_the_content_and_admits_it_is_degraded():
    result = summarizer._fallback("One thing happened. Then another thing.")
    assert result["degraded"] is True
    assert result["key_points"]
    # Must never invent decisions nobody made.
    assert result["decisions"] == []
    assert result["actions"] == []


def test_long_transcripts_are_trimmed_from_the_middle():
    long_text = "A" * 20_000 + "ZZZEND"
    trimmed = summarizer._trim(long_text)
    assert len(trimmed) < len(long_text)
    assert trimmed.endswith("ZZZEND")      # the closing minutes are kept
    assert "omitted" in trimmed


def test_string_actions_are_normalised_to_task_owner():
    parsed = summarizer._as_list(["do the thing"])
    assert parsed == ["do the thing"]


# ── Word export ──────────────────────────────────────────────────────────
@pytest.mark.skipif(not export_docx.is_available(), reason="python-docx not installed")
def test_word_document_contains_the_minutes(tmp_path):
    from docx import Document

    summary = {
        "summary": "Planning call.",
        "key_points": ["Rollout schedule"],
        "decisions": ["Postpone Duisburg"],
        "actions": [{"task": "Prepare measurements", "owner": "Hocine"}],
        "open_questions": ["Start in September?"],
        "degraded": False,
    }
    path = export_docx.write_minutes(_state(), summary, out_dir=tmp_path)
    assert path.is_file() and path.suffix == ".docx"

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Planning call." in text
    assert "Rollout schedule" in text
    assert "Postpone Duisburg" in text
    # The full transcript must be there, not just the summary.
    assert "We need to decide the rollout." in text
    # Action items go in a table with an owner column.
    owners = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "Hocine" in owners


@pytest.mark.skipif(not export_docx.is_available(), reason="python-docx not installed")
def test_degraded_summary_is_flagged_in_the_document(tmp_path):
    from docx import Document

    path = export_docx.write_minutes(
        _state(), summarizer._fallback("Something was said."), out_dir=tmp_path
    )
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "could not be produced" in text


@pytest.mark.parametrize("title,expected_start", [
    ("Fibre rollout planning", "Fibre-rollout-planning"),
    ("Q3 / budget: review!", "Q3-budget-review"),
    ("", "20260811-140000"),
])
def test_titles_become_safe_filenames(title, expected_start):
    """Titles come from speech and can contain anything."""
    assert export_docx._safe_stem(title, "20260811-140000").startswith(expected_start)


def test_filename_never_escapes_the_directory():
    assert "/" not in export_docx._safe_stem("../../etc/passwd", "fallback")
    assert ".." not in export_docx._safe_stem("../../etc/passwd", "fallback")
