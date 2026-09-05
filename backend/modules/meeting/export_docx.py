"""Write meeting minutes to a Word (.docx) document."""
from __future__ import annotations

import logging
import re
from pathlib import Path

from backend.core.config import config

log = logging.getLogger("plasma.meeting")

MINUTES_DIR: Path = config.PLASMA_DIR / "meetings"


def _safe_stem(title: str, fallback: str) -> str:
    """Filename-safe stem — the title comes from speech and can contain anything."""
    cleaned = re.sub(r"[^\w\s\-]", "", title, flags=re.UNICODE).strip()
    cleaned = re.sub(r"\s+", "-", cleaned)
    return (cleaned[:60] or fallback).strip("-") or fallback


def is_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def write_minutes(state, summary: dict, out_dir: Path | None = None) -> Path:
    """Render minutes to .docx and return the path.

    Raises RuntimeError if python-docx isn't installed, so the caller can tell
    the user something actionable instead of failing silently.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed — run: pip install python-docx"
        ) from e

    out_dir = Path(out_dir) if out_dir else MINUTES_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(state.title, level=0)

    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run(state.started_at.strftime("%d.%m.%Y %H:%M"))
    meta.add_run("    Duration: ").bold = True
    meta.add_run(f"{state.duration_min:.0f} min")

    if summary.get("degraded"):
        note = doc.add_paragraph()
        run = note.add_run(
            "Note: an automatic summary could not be produced for this meeting. "
            "The full transcript is included below."
        )
        run.italic = True

    if summary.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary["summary"])

    def bullets(heading: str, items: list[str]) -> None:
        if not items:
            return
        doc.add_heading(heading, level=1)
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")

    bullets("Key points", summary.get("key_points") or [])
    bullets("Decisions", summary.get("decisions") or [])

    actions = summary.get("actions") or []
    if actions:
        doc.add_heading("Action items", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Task"
        hdr[1].text = "Owner"
        for a in actions:
            row = table.add_row().cells
            row[0].text = str(a.get("task", ""))
            row[1].text = str(a.get("owner", "")) or "—"

    bullets("Open questions", summary.get("open_questions") or [])

    doc.add_page_break()
    doc.add_heading("Full transcript", level=1)
    for seg in state.segments:
        p = doc.add_paragraph()
        stamp = str(seg.get("ts", ""))[11:16]      # HH:MM
        if stamp:
            r = p.add_run(f"[{stamp}] ")
            r.bold = True
            r.font.size = Pt(9)
        p.add_run(seg.get("text", ""))

    path = out_dir / f"{_safe_stem(state.title, state.meeting_id)}.docx"
    doc.save(path)
    log.info("Meeting minutes written: %s", path)
    return path
